from __future__ import annotations

import re
from pathlib import Path

from docx import Document


REPORTS_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = REPORTS_DIR / "docx"


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def _clean_inline_markdown(text: str) -> str:
    # Bold/italic (keep contents)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)

    # Inline code
    text = re.sub(r"`(.+?)`", r"\1", text)

    # Links: [label](url) -> label (url)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)

    return text


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(_clean_inline_markdown(text))
    if style:
        p.style = style


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    in_code_block = False

    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip("\n").rstrip("\r").rstrip()

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            # Keep code blocks as plain paragraphs.
            _add_paragraph(doc, line)
            continue

        if not line.strip():
            # Single blank line as spacing.
            doc.add_paragraph("")
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            title = _clean_inline_markdown(title)
            # Clamp heading level to a reasonable range for Word.
            level = min(level, 4)
            doc.add_heading(title, level=level)
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            _add_paragraph(doc, numbered_match.group(2).strip(), style="List Number")
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            _add_paragraph(doc, stripped[2:].strip(), style="List Bullet")
            continue

        _add_paragraph(doc, line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    md_files = sorted(REPORTS_DIR.glob("week*_*.md"))
    if not md_files:
        raise SystemExit("No weekly markdown reports found to convert.")

    for md_path in md_files:
        docx_path = OUTPUT_DIR / f"{md_path.stem}.docx"
        convert_markdown_to_docx(md_path, docx_path)
        print(f"Wrote: {docx_path}")


if __name__ == "__main__":
    main()
